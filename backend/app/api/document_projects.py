"""
文档项目相关 API 路由
"""
from fastapi import APIRouter, HTTPException, Query, Request
from fastapi.responses import StreamingResponse, HTMLResponse, Response
from pydantic import BaseModel, Field
from typing import List, Optional
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn  # 用于设置字体
from docx.oxml import OxmlElement
import io
import re
from urllib.parse import quote
import markdown
from bs4 import BeautifulSoup
import logging

from app.models.document_project import document_project_storage
from app.services.document_generator import document_generator_service

logger = logging.getLogger(__name__)


def convert_quotes_to_chinese(text):
    """
    将英文引号转换为中文引号
    如果引号内的文本包含中文字符，使用中文引号
    """
    def replace_double_quotes(match):
        content = match.group(1)
        # 如果内容包含中文字符，使用中文引号
        if any('\u4e00' <= char <= '\u9fff' for char in content):
            return '\u201c' + content + '\u201d'  # 中文引号 U+201C U+201D
        else:
            return match.group(0)  # 保持原样（英文引号）

    def replace_single_quotes(match):
        content = match.group(1)
        # 如果内容包含中文字符，使用中文引号
        if any('\u4e00' <= char <= '\u9fff' for char in content):
            return '\u2018' + content + '\u2019'  # 中文引号 U+2018 U+2019
        else:
            return match.group(0)  # 保持原样（英文引号）

    # 处理双引号 "..."
    text = re.sub(r'"([^"]+)"', replace_double_quotes, text)
    # 处理单引号 '...'
    text = re.sub(r"'([^']+)'", replace_single_quotes, text)

    return text


def add_markdown_to_paragraph(paragraph, md_text):
    """
    将 Markdown 文本添加到段落，保留格式（粗体、斜体、代码、公式等）
    中文引号使用宋体，英文内容使用Times New Roman
    """
    # 先将英文引号转换为中文引号
    md_text = convert_quotes_to_chinese(md_text)

    # 使用占位符保存格式化文本
    placeholders = {}
    placeholder_count = [0]

    def get_placeholder():
        placeholder_count[0] += 1
        return f"__PLACEHOLDER_{placeholder_count[0]}__"

    # 第零步：处理数学公式（最优先处理，避免被其他规则干扰）
    def replace_math_block(text):
        # 块级公式 $$...$$ (使用DOTALL匹配跨行)
        pattern = r'\$\$([^\$]+?)\$\$'
        def replacement(match):
            key = get_placeholder()
            placeholders[key] = (match.group(1), 'math-block')
            return key
        return re.sub(pattern, replacement, text, flags=re.DOTALL)

    def replace_math_inline(text):
        # 行内公式 $...$ (必须在同一行内，避免跨块级公式匹配)
        # 匹配同一行内的$...$模式
        pattern = r'([^\n$]+?)\$'  # 这个模式会先匹配非$的字符，然后是$，但实际上我们需要更精确的模式
        # 修正：使用负向前瞻确保匹配的不是$$的一部分
        # 更好的方法：按行处理，每行单独匹配$...$
        lines = text.split('\n')
        result_lines = []
        for line in lines:
            # 对每一行，替换$...$为占位符
            def replace_inline_in_line(match):
                key = get_placeholder()
                placeholders[key] = (match.group(1), 'math-inline')
                return key
            line = re.sub(r'\$([^$]+?)\$', replace_inline_in_line, line)
            result_lines.append(line)
        return '\n'.join(result_lines)

    # 第一步：处理粗体 **text**
    def replace_bold(text):
        pattern = r'\*\*(.+?)\*\*'
        def replacement(match):
            key = get_placeholder()
            placeholders[key] = (match.group(1), 'bold')
            return key
        return re.sub(pattern, replacement, text)

    # 第二步：处理斜体 *text*（避免匹配 **）
    def replace_italic(text):
        pattern = r'(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)'
        def replacement(match):
            key = get_placeholder()
            placeholders[key] = (match.group(1), 'italic')
            return key
        return re.sub(pattern, replacement, text)

    # 第三步：处理代码 `text`
    def replace_code(text):
        pattern = r'`(.+?)`'
        def replacement(match):
            key = get_placeholder()
            placeholders[key] = (match.group(1), 'code')
            return key
        return re.sub(pattern, replacement, text)

    # 按顺序替换（公式优先）
    result = replace_math_block(md_text)  # 块级公式优先
    result = replace_math_inline(result)  # 然后行内公式
    result = replace_bold(result)
    result = replace_italic(result)
    result = replace_code(result)

    # 还原并添加到段落
    parts = re.split(r'(__PLACEHOLDER_\d+__)', result)

    # 中文引号字符列表（使用Unicode编码）
    chinese_quotes = ['\u201c', '\u201d', '\u2018', '\u2019', '\u300c', '\u300d', '\u300e', '\u300f']

    for part in parts:
        if part in placeholders:
            text, fmt_type = placeholders[part]

            # 特殊处理：数学公式
            if fmt_type in ['math-inline', 'math-block']:
                # 强制使用 latex2word 渲染所有公式
                try:
                    from latex2word import LatexToWordElement

                    # Debug: 记录LaTeX内容
                    logger.debug(f"处理公式: {fmt_type}")
                    logger.debug(f"LaTeX (repr): {repr(text)}")
                    logger.debug(f"LaTeX长度: {len(text)}")

                    # 创建 LaTeX 到 Word 的转换对象
                    latex_to_word = LatexToWordElement(text)

                    if fmt_type == 'math-block':
                        # 块级公式：创建新段落并居中
                        from docx.oxml import OxmlElement

                        para_element = paragraph._element
                        body = para_element.getparent()
                        if body is None:
                            latex_to_word.add_latex_to_paragraph(paragraph)
                        else:
                            # 创建新段落元素
                            new_para_element = OxmlElement('w:p')
                            body.insert(body.index(para_element) + 1, new_para_element)

                            # 创建段落对象并添加公式
                            from docx.text.paragraph import Paragraph
                            formula_para = Paragraph(new_para_element, paragraph._parent)
                            formula_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                            # 设置段落间距，避免多余的空行
                            formula_para.paragraph_format.space_before = Pt(0)
                            formula_para.paragraph_format.space_after = Pt(0)
                            formula_para.paragraph_format.line_spacing = 1.5

                            latex_to_word.add_latex_to_paragraph(formula_para)
                    else:
                        # 行内公式：直接插入当前段落
                        latex_to_word.add_latex_to_paragraph(paragraph)
                        logger.debug(f"行内公式已添加到段落")

                except Exception as e:
                    # 如果 latex2word 失败，记录错误
                    import traceback
                    logger.error(f"LaTeX 公式转换失败: {e}\nLaTeX: {text[:50]}...\nTraceback: {traceback.format_exc()}")
                continue

            # 普通格式化文本
            for char in text:
                run = paragraph.add_run(char)

                # 如果是中文引号，设置为宋体
                if char in chinese_quotes:
                    run.font.name = '宋体'
                    run.font.size = Pt(12)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                else:
                    # 普通文本设置格式
                    if fmt_type == 'bold':
                        run.bold = True
                    elif fmt_type == 'italic':
                        run.italic = True
                    elif fmt_type == 'code':
                        run.font.name = 'Courier New'
                        run.font.size = Pt(10)

        else:
            # 普通文本，逐个字符处理
            if part:
                for char in part:
                    run = paragraph.add_run(char)

                    # 如果是中文引号，设置为宋体
                    if char in chinese_quotes:
                        run.font.name = '宋体'
                        run.font.size = Pt(12)
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')




router = APIRouter()


class CreateProjectRequest(BaseModel):
    """创建项目请求"""
    title: str = Field(..., description="项目标题")
    folderIds: List[str] = Field(default=[], description="知识库ID列表")
    outline: Optional[List[dict]] = Field(None, description="大纲树形结构")
    content: Optional[List[dict]] = Field(None, description="章节内容")


class UpdateOutlineRequest(BaseModel):
    """更新大纲请求"""
    outline: List[dict] = Field(..., description="大纲树形结构")
    locked: bool = Field(False, description="是否锁定大纲")


class GenerateOutlineRequest(BaseModel):
    """生成大纲请求"""
    topic: str = Field(..., description="研究主题")


@router.post("")
async def create_project(request: CreateProjectRequest):
    """
    创建文档项目

    创建新的文档生成项目，可以选择多个知识库，也可以直接传入大纲和内容
    """
    try:
        project = document_project_storage.create_project(
            title=request.title,
            folder_ids=request.folderIds,
            outline=request.outline,
            content=request.content,
        )
        return project

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"创建项目失败: {str(e)}")


@router.get("")
async def list_projects(
    skip: int = Query(0, description="跳过数量"),
    limit: int = Query(20, description="返回数量")
):
    """
    获取项目列表

    获取所有文档项目
    """
    try:
        projects, total = document_project_storage.list_projects(
            skip=skip,
            limit=limit,
        )

        return {
            "projects": projects,
            "total": total
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"获取项目列表失败: {str(e)}")


@router.get("/{project_id}")
async def get_project(project_id: str):
    """
    获取项目详情

    获取项目的完整信息，包括大纲和内容
    """
    try:
        project = document_project_storage.get_project(project_id)
        if not project:
            raise HTTPException(status_code=404, detail="项目不存在")

        return project

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"获取项目失败: {str(e)}")


@router.put("/{project_id}/outline")
async def update_outline(project_id: str, request: UpdateOutlineRequest):
    """
    更新大纲

    更新项目的大纲并设置锁定状态
    """
    try:
        project = document_project_storage.get_project(project_id)
        if not project:
            raise HTTPException(status_code=404, detail="项目不存在")

        current_locked = project.get("outlineLocked", False)

        # 只有在当前已锁定且请求保持锁定状态时才拒绝修改
        # 允许解锁操作（locked=False）或首次锁定
        if current_locked and request.locked:
            raise HTTPException(status_code=400, detail="大纲已锁定，无法修改")

        updated_project = document_project_storage.update_outline(
            project_id=project_id,
            outline=request.outline,
            locked=request.locked,
        )

        return updated_project

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"更新大纲失败: {str(e)}")


@router.post("/{project_id}/generate-outline")
async def generate_outline(project_id: str, request: GenerateOutlineRequest):
    """
    生成大纲

    基于研究主题和知识库内容自动生成文档大纲
    """
    try:
        project = document_project_storage.get_project(project_id)
        if not project:
            raise HTTPException(status_code=404, detail="项目不存在")

        # 检查大纲是否已锁定
        if project.get("outlineLocked", False):
            raise HTTPException(status_code=400, detail="大纲已锁定，无法重新生成")

        # 生成大纲
        outline = document_generator_service.generate_outline(
            topic=request.topic,
            folder_ids=project.get("folderIds", []),
        )

        # 更新项目
        updated_project = document_project_storage.update_outline(
            project_id=project_id,
            outline=outline,
            locked=False,  # 生成后不锁定，需要用户确认
        )

        return updated_project

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"生成大纲失败: {str(e)}")


@router.delete("/{project_id}")
async def delete_project(project_id: str):
    """
    删除项目

    删除文档项目
    """
    try:
        success = document_project_storage.delete_project(project_id)
        if not success:
            raise HTTPException(status_code=404, detail="项目不存在")

        return {"message": "删除成功"}

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"删除项目失败: {str(e)}")


class GenerateContentRequest(BaseModel):
    """生成内容请求"""
    sectionId: str = Field(..., description="章节ID")
    sectionTitle: str = Field(..., description="章节标题")
    contextSections: List[str] = Field(default=[], description="上下文章节路径")
    customPrompt: str = Field(default="", description="自定义生成需求")


class RegenerateParagraphRequest(BaseModel):
    """重新生成段落请求"""
    sectionId: str = Field(..., description="章节ID")
    sectionTitle: str = Field(..., description="章节标题")
    contextSections: List[str] = Field(default=[], description="上下文章节路径")
    customPrompt: str = Field(default="", description="自定义生成需求")


class UpdateParagraphRequest(BaseModel):
    """更新段落请求"""
    sectionId: str = Field(..., description="章节ID")
    paragraphId: str = Field(..., description="段落ID")
    content: str = Field(..., description="段落内容")


@router.post("/{project_id}/generate-content")
async def generate_section_content(project_id: str, request: GenerateContentRequest):
    """
    生成章节内容

    为指定章节生成内容，基于RAG检索的相关内容
    """
    try:
        project = document_project_storage.get_project(project_id)
        if not project:
            raise HTTPException(status_code=404, detail="项目不存在")

        # 获取所有文档ID
        document_ids = []
        for folder_id in project.get("folderIds", []):
            from app.models.document import storage as doc_storage
            documents, _ = doc_storage.list_documents(folder=folder_id)
            document_ids.extend([doc["id"] for doc in documents])

        # 获取完整大纲（用于上下文）
        outline = project.get("outline", [])

        # 生成内容
        result = document_generator_service.generate_section_content(
            section_title=request.sectionTitle,
            section_id=request.sectionId,
            document_ids=document_ids,
            context_sections=request.contextSections if request.contextSections else None,
            custom_prompt=request.customPrompt if request.customPrompt else None,
            full_outline=outline  # 传递完整大纲
        )

        # 保存到项目
        document_project_storage.add_section_content(
            project_id=project_id,
            section_id=request.sectionId,
            content={
                "sectionId": request.sectionId,
                "paragraphs": [result],
                "sources": result.get("sources", [])
            }
        )

        return result

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"生成内容失败: {str(e)}")


@router.post("/{project_id}/regenerate-paragraph")
async def regenerate_paragraph(project_id: str, request: RegenerateParagraphRequest):
    """
    重新生成段落

    重新生成指定章节的段落内容，保存旧版本到versions中
    """
    try:
        project = document_project_storage.get_project(project_id)
        if not project:
            raise HTTPException(status_code=404, detail="项目不存在")

        # 获取所有文档ID
        document_ids = []
        for folder_id in project.get("folderIds", []):
            from app.models.document import storage as doc_storage
            documents, _ = doc_storage.list_documents(folder=folder_id)
            document_ids.extend([doc["id"] for doc in documents])

        # 获取当前段落（用于保存版本）
        sections = project.get("sections", {})
        section = sections.get(request.sectionId, {})
        paragraphs = section.get("paragraphs", [])

        # 准备当前段落信息和版本历史
        current_paragraph = None
        existing_versions = []

        if paragraphs:
            # 获取最新的段落
            current_paragraph = paragraphs[-1]
            # 获取已有的历史版本
            existing_versions = current_paragraph.get("versions", [])

        # 获取完整大纲（用于上下文）
        outline = project.get("outline", [])

        # 重新生成
        new_paragraph_data = document_generator_service.regenerate_paragraph(
            section_title=request.sectionTitle,
            section_id=request.sectionId,
            document_ids=document_ids,
            context_sections=request.contextSections if request.contextSections else None,
            custom_prompt=request.customPrompt if request.customPrompt else None,
            full_outline=outline  # 传递完整大纲
        )

        # 如果有当前段落，保存到版本历史
        if current_paragraph:
            # 将当前段落添加到版本历史
            existing_versions.append({
                "content": current_paragraph.get("content", ""),
                "timestamp": current_paragraph.get("timestamp", ""),
                "sources": current_paragraph.get("sources", [])
            })

        # 创建新段落（包含历史版本）
        updated_paragraph = {
            "paragraph_id": new_paragraph_data.get("paragraph_id"),
            "section_id": new_paragraph_data.get("section_id"),
            "content": new_paragraph_data.get("content"),
            "sources": new_paragraph_data.get("sources", []),
            "timestamp": new_paragraph_data.get("timestamp"),
            "versions": existing_versions  # 保存所有历史版本
        }

        # 更新项目（只保留一个段落，历史在versions中）
        document_project_storage.add_section_content(
            project_id=project_id,
            section_id=request.sectionId,
            content={
                "sectionId": request.sectionId,
                "paragraphs": [updated_paragraph],  # 只保留当前段落
                "sources": updated_paragraph.get("sources", [])
            }
        )

        return updated_paragraph

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"重新生成失败: {str(e)}")


@router.put("/{project_id}/paragraph")
async def update_paragraph(project_id: str, request: UpdateParagraphRequest):
    """
    更新段落内容（编辑）

    直接修改段落内容，不保存版本
    """
    try:
        updated_project = document_project_storage.update_paragraph(
            project_id=project_id,
            section_id=request.sectionId,
            paragraph_id=request.paragraphId,
            content=request.content,
            save_version=False  # 编辑时不保存版本
        )

        if not updated_project:
            raise HTTPException(status_code=404, detail="段落不存在")

        return {"message": "更新成功"}

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"更新失败: {str(e)}")


class RestoreParagraphVersionRequest(BaseModel):
    """恢复段落版本请求"""
    sectionId: str = Field(..., description="章节ID")
    paragraphId: str = Field(..., description="段落ID")
    versionIndex: int = Field(..., description="版本索引")


@router.post("/{project_id}/restore-paragraph-version")
async def restore_paragraph_version(project_id: str, request: RestoreParagraphVersionRequest):
    """
    恢复段落到指定版本

    将段落恢复到历史版本
    """
    try:
        updated_project = document_project_storage.restore_paragraph_version(
            project_id=project_id,
            section_id=request.sectionId,
            paragraph_id=request.paragraphId,
            version_index=request.versionIndex
        )

        if not updated_project:
            raise HTTPException(status_code=404, detail="段落或版本不存在")

        return updated_project

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"恢复失败: {str(e)}")


@router.get("/{project_id}/export-word")
async def export_word(
    project_id: str,
    title: str = Query(default="文档", description="导出文件名")
):
    """
    导出为 Word 文档

    将项目的大纲和内容导出为 Word 文档

    Args:
        project_id: 项目ID
        title: 导出文件名（不含扩展名）
    """
    try:
        project = document_project_storage.get_project(project_id)
        if not project:
            raise HTTPException(status_code=404, detail="项目不存在")

        # 创建 Word 文档
        doc = Document()

        # 小四号字体 = 12磅
        font_size_small4 = Pt(12)

        # 设置文档默认样式（Normal 样式）
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = font_size_small4
        # 设置中文字体
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # 设置文档标题
        title = project.get("title", "文档")
        title_heading = doc.add_heading(title, 0)
        title_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 设置标题格式
        for run in title_heading.runs:
            run.font.size = font_size_small4
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.underline = False
            # 使用 qn 设置字体
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # 移除段落边框
        try:
            pPr = title_heading._element.get_or_add_pPr()
            pBdr = pPr.find(qn('w:pBdr'))
            if pBdr is not None:
                pPr.remove(pBdr)
        except:
            pass

        # 设置标题段后间距
        title_heading.paragraph_format.line_spacing = 1.5
        title_heading.paragraph_format.space_before = Pt(0)
        title_heading.paragraph_format.space_after = Pt(0)

        # 添加大纲内容
        outline = project.get("outline", [])
        sections = project.get("sections", {})

        def add_outline_to_doc(nodes: List, doc_level: int = 1):
            """递归添加大纲节点到 Word 文档"""
            for node in nodes:
                node_id = node.get("id")
                label = node.get("label", "无标题")

                # 添加章节标题
                heading = doc.add_heading(label, level=min(doc_level, 9))

                # 设置标题格式
                for run in heading.runs:
                    run.font.size = font_size_small4
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    run.font.underline = False
                    # 使用 qn 设置字体
                    run.font.name = 'Times New Roman'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

                # 移除段落边框
                try:
                    pPr = heading._element.get_or_add_pPr()
                    pBdr = pPr.find(qn('w:pBdr'))
                    if pBdr is not None:
                        pPr.remove(pBdr)
                except:
                    pass

                # 设置标题段后间距和行距
                heading.paragraph_format.line_spacing = 1.5
                heading.paragraph_format.space_before = Pt(0)
                heading.paragraph_format.space_after = Pt(0)

                # 添加章节内容（如果有）
                if node_id and node_id in sections:
                    section_data = sections[node_id]
                    paragraphs = section_data.get("paragraphs", [])

                    if paragraphs:
                        # 只显示最新段落（当前版本）
                        current_para = paragraphs[-1]
                        content = current_para.get("content", "")

                        # 调试：打印内容的前200个字符
                        print(f"\n[DEBUG] 导出章节: {label}")
                        print(f"[DEBUG] 内容前200字符: {repr(content[:200])}")
                        print(f"[DEBUG] 内容长度: {len(content)}")
                        print(f"[DEBUG] 包含 ** : {'**' in content}")
                        print(f"[DEBUG] 包含 * : {'*' in content}")

                        if content.strip():
                            # 直接按双换行符分割内容，每个逻辑段落单独处理
                            import re
                            paragraphs_text = re.split(r'\n\s*\n', content)

                            for para_text in paragraphs_text:
                                para_text = para_text.strip()
                                if not para_text:
                                    continue

                                # 添加段落并处理 Markdown 格式
                                para = doc.add_paragraph()

                                # 设置段落格式
                                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                para.paragraph_format.first_line_indent = Inches(0.32)  # 首行缩进
                                para.paragraph_format.line_spacing = 1.5  # 1.5倍行距
                                para.paragraph_format.space_before = Pt(0)
                                para.paragraph_format.space_after = Pt(0)

                                # 将 Markdown 内容添加到段落（保留格式）
                                # add_markdown_to_paragraph函数内部会处理所有公式（包括行内和块级）
                                add_markdown_to_paragraph(para, para_text)

                                # 设置所有 run 的字体（跳过已经是宋体的中文引号和代码块）
                                for run in para.runs:
                                    # 跳过代码块和已经是宋体的中文引号
                                    if run.font.name == 'Courier New':
                                        continue
                                    if run.font.name == '宋体':
                                        # 中文引号已经是宋体，只需设置字体大小
                                        run.font.size = font_size_small4
                                        continue

                                    # 其他文本设置为Times New Roman
                                    run.font.size = font_size_small4
                                    run.font.name = 'Times New Roman'
                                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

                # 递归处理子节点
                children = node.get("children", [])
                if children:
                    add_outline_to_doc(children, doc_level + 1)

        # 添加大纲内容
        if outline:
            add_outline_to_doc(outline)

        # 保存到内存
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        # 生成文件名（使用传递的 title 参数）
        filename = f"{title}.docx"
        encoded_filename = quote(filename)

        # 返回文件流
        return StreamingResponse(
            file_stream,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"
            }
        )

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"导出失败: {str(e)}")


@router.get("/{project_id}/preview-html")
async def preview_html(project_id: str, request: Request):
    """
    生成HTML预览

    生成项目内容的HTML预览，用于PDF导出
    """
    try:
        # 获取后端基础 URL
        backend_base_url = f"{request.url.scheme}://{request.url.netloc}"

        project = document_project_storage.get_project(project_id)
        if not project:
            raise HTTPException(status_code=404, detail="项目不存在")

        title = project.get("title", "文档")
        outline = project.get("outline", [])
        sections = project.get("sections", {})

        # 用于存储提取的公式
        extracted_formulas = []
        formula_counter = [0]  # 使用列表以便在嵌套函数中修改

        def protect_formulas(content: str) -> str:
            """保护 LaTeX 公式，防止被 Markdown 处理"""
            # 先提取块级公式
            content = re.sub(
                r'\$\$([^\$]+?)\$\$',
                lambda m: _extract_formula(m.group(1), 'block'),
                content,
                flags=re.DOTALL
            )
            # 再提取行内公式
            content = re.sub(
                r'\$([^\$\n]+?)\$',
                lambda m: _extract_formula(m.group(1), 'inline'),
                content
            )
            return content

        def _extract_formula(latex: str, fmt_type: str) -> str:
            """提取公式并返回占位符"""
            idx = formula_counter[0]
            formula_counter[0] += 1
            extracted_formulas.append({
                'index': idx,
                'latex': latex.strip(),
                'type': fmt_type
            })
            # 使用特殊占位符，Markdown 不会处理
            return f"MATHFORMULA{idx}PLACEHOLDER"

        # 生成HTML
        html_content = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{title}</title>
    <!-- KaTeX CSS - 先加载 KaTeX 样式 -->
    <link rel="stylesheet" href="{backend_base_url}/static/katex/katex.min.css">
    <style>
        * {{
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }}

        @page {{
            size: A4;
            margin: 20mm;
        }}

        html, body {{
            margin: 0;
            padding: 0;
            width: 100%;
            height: 100%;
        }}

        body {{
            font-family: 'Times New Roman', '宋体', serif;
            font-size: 12pt;
            line-height: 1.5;
            max-width: 210mm;
            margin: 20mm auto;
            padding: 20mm;
            background: white;
        }}

        /* 中文引号使用宋体 */
        .chinese-quote {{
            font-family: '宋体', serif;
        }}

        h1, h2, h3 {{
            font-size: 12pt;
            font-weight: bold;
            margin: 0;
            margin-top: 0;
            margin-bottom: 0;
            padding: 0;
            padding-top: 0;
            padding-bottom: 0;
            color: #000;
            line-height: 1.5;
            text-align: left;
        }}

        h1 {{
            text-align: center;
        }}

        p {{
            text-align: justify;
            text-indent: 2em;
            margin: 0;
            margin-top: 0;
            margin-bottom: 0;
            padding: 0;
            padding-top: 0;
            padding-bottom: 0;
            line-height: 1.5;
            font-size: 12pt;
            page-break-inside: avoid;
        }}

        @media print {{
            body {{
                margin: 0;
                padding: 20mm;
            }}

            h1, h2, h3, p {{
                margin: 0 !important;
                padding: 0 !important;
            }}
        }}

        /* KaTeX 公式样式 - 完全重置间距，使用最高优先级 */
        html body .katex-display,
        html body div.katex-display,
        body > .katex-display {{
            margin: 0 !important;
            padding: 0 !important;
            display: block !important;
            line-height: 1.5 !important;
        }}

        html body .katex-display > .katex,
        html body .katex {{
            line-height: 1.5 !important;
        }}

        html body .katex-display > .katex,
        html body .katex-display > .katex-display {{
            display: inline-block;
            margin: 0 !important;
            padding: 0 !important;
        }}

        /* 确保公式容器没有额外间距 */
        html body .katex-display > .katex > .katex-html,
        html body .katex-display > .katex > .katex-html > .base {{
            margin: 0 !important;
            padding: 0 !important;
        }}
    </style>
</head>
<body>
    <h1>{title}</h1>
"""

        def add_outline_html(nodes: List, level: int = 1) -> str:
            """递归生成大纲HTML"""
            html = ""
            for node in nodes:
                node_id = node.get("id")
                label = node.get("label", "无标题")

                # 根据层级选择标题标签
                if level == 1:
                    html += f"<h2>{label}</h2>"
                elif level == 2:
                    html += f"<h3>{label}</h3>"
                else:
                    html += f"<h3>{label}</h3>"

                # 添加章节内容
                if node_id and node_id in sections:
                    section_data = sections[node_id]
                    paragraphs = section_data.get("paragraphs", [])

                    if paragraphs:
                        current_para = paragraphs[-1]
                        content = current_para.get("content", "")

                        if content.strip():
                            # 先将英文引号转换为中文引号
                            content = convert_quotes_to_chinese(content)

                            # 然后将中文引号包裹在span中（使用Unicode编码）
                            content = re.sub(
                                r'([\u201c\u201d\u2018\u2019\u300c\u300d\u300e\u300f])',
                                r'<span class="chinese-quote">\1</span>',
                                content
                            )

                            # 保护 LaTeX 公式，防止被 Markdown 处理
                            content = protect_formulas(content)

                            # 将 Markdown 转换为 HTML
                            md_html = markdown.markdown(
                                content,
                                extensions=[
                                    'extra',          # 额外功能（表格、列表等）
                                    'nl2br',          # 换行符转 <br>
                                    'sane_lists',     # 更好的列表支持
                                ]
                            )

                            # 将占位符替换回 LaTeX 公式（供 KaTeX auto-render 处理）
                            for formula in extracted_formulas:
                                idx = formula['index']
                                latex = formula['latex']
                                fmt_type = formula['type']
                                if fmt_type == 'block':
                                    placeholder = f"MATHFORMULA{idx}PLACEHOLDER"
                                    replacement = f"$${latex}$$"
                                else:
                                    placeholder = f"MATHFORMULA{idx}PLACEHOLDER"
                                    replacement = f"${latex}$"
                                md_html = md_html.replace(placeholder, replacement)

                            html += md_html

                # 递归处理子节点
                children = node.get("children", [])
                if children:
                    html += add_outline_html(children, level + 1)

            return html

        if outline:
            html_content += add_outline_html(outline)

        html_content += f"""    <script>
        // 动态加载 KaTeX 资源（确保加载顺序）
        (function() {{
            const backendUrl = '{backend_base_url}';

            // 加载 KaTeX 核心库
            const katexScript = document.createElement('script');
            katexScript.src = backendUrl + '/static/katex/katex.min.js';
            katexScript.onload = function() {{
                console.log('✓ KaTeX 核心库已加载');

                // KaTeX 核心库加载完成后，再加载 auto-render 插件
                const autoRenderScript = document.createElement('script');
                autoRenderScript.src = backendUrl + '/static/katex/contrib/auto-render.min.js';
                autoRenderScript.onload = function() {{
                    console.log('✓ KaTeX auto-render 已加载');

                    // auto-render 加载完成后，立即渲染公式
                    try {{
                        renderMathInElement(document.body, {{
                            delimiters: [
                                {{left: '$$', right: '$$', display: true}},
                                {{left: '$', right: '$', display: false}}
                            ],
                            throwOnError: false
                        }});
                        console.log('✓ KaTeX 公式渲染完成');
                    }} catch (e) {{
                        console.error('KaTeX 渲染失败:', e);
                    }}

                    // 渲染完成后延迟触发打印
                    setTimeout(function() {{
                        window.print();
                    }}, 1000);
                }};
                autoRenderScript.onerror = function() {{
                    console.error('✗ KaTeX auto-render 加载失败');
                }};
                document.head.appendChild(autoRenderScript);
            }};
            katexScript.onerror = function() {{
                console.error('✗ KaTeX 核心库加载失败，请检查后端静态文件服务');
            }};
            document.head.appendChild(katexScript);
        }})();
    </script>
</body>
</html>"""

        # 使用 Response 而不是 HTMLResponse，避免 Content-Length 计算错误
        return Response(
            content=html_content,
            media_type="text/html; charset=utf-8"
        )

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"预览生成失败: {str(e)}")


